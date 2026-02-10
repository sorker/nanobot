#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件格式转换工具
支持多种文件格式之间的转换
"""

import os
import tempfile
from typing import Optional, Dict, Any
from pathlib import Path
import markdown
from markdown.extensions import tables, codehilite, fenced_code, toc
import pypandoc
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import weasyprint
from bs4 import BeautifulSoup
import requests
from io import BytesIO
import asyncio
import base64
import re
from loguru import logger

# 可选导入playwright
try:
    from playwright.async_api import async_playwright
    PLAYWRIGHT_AVAILABLE = True
except ImportError:
    PLAYWRIGHT_AVAILABLE = False
    # 注意：此时logger还未定义，使用print输出警告
    import sys
    print("警告: Playwright未安装，PDF转换将使用weasyprint作为备用方案", file=sys.stderr)


# ============== 模块级浏览器实例管理器（单例模式） ==============
class BrowserManager:
    """
    模块级浏览器实例管理器
    确保整个Python进程中只有一个浏览器实例，避免重复启动
    """
    _instance = None
    
    def __init__(self):
        self.playwright = None
        self.browser = None
        self.is_initialized = False
        self._init_lock = None  # 延迟创建Lock
    
    @classmethod
    def get_instance(cls):
        """获取单例实例"""
        if cls._instance is None:
            cls._instance = cls()
        return cls._instance
    
    async def initialize(self):
        """初始化浏览器（延迟初始化）"""
        if not PLAYWRIGHT_AVAILABLE:
            raise RuntimeError("Playwright未安装，无法初始化浏览器")
        
        if self.is_initialized and self.browser:
            return
        
        # 延迟创建Lock（在异步上下文中）
        if self._init_lock is None:
            self._init_lock = asyncio.Lock()
        
        async with self._init_lock:
            # 双重检查
            if self.is_initialized and self.browser:
                return
            
            try:
                logger.info("🚀 正在初始化模块级浏览器实例...")
                self.playwright = await async_playwright().start()
                self.browser = await self.playwright.chromium.launch(
                    headless=True,
                    args=['--no-sandbox', '--disable-setuid-sandbox']
                )
                self.is_initialized = True
                logger.info("✅ 模块级浏览器实例初始化成功（将被重用以提升性能）")
            except Exception as e:
                logger.error(f"❌ 浏览器初始化失败: {e}")
                self.is_initialized = False
                raise
    
    async def get_browser(self):
        """获取浏览器实例，如果未初始化则自动初始化"""
        if not self.is_initialized or not self.browser:
            await self.initialize()
        return self.browser
    
    async def close(self):
        """关闭浏览器实例"""
        if self.browser:
            try:
                await self.browser.close()
                logger.info("浏览器实例已关闭")
            except Exception as e:
                logger.warning(f"关闭浏览器失败: {e}")
            finally:
                self.browser = None
        
        if self.playwright:
            try:
                await self.playwright.stop()
            except Exception as e:
                logger.warning(f"停止playwright失败: {e}")
            finally:
                self.playwright = None
        
        self.is_initialized = False
    
    def __del__(self):
        """析构时清理资源"""
        if self.is_initialized:
            # 注意：这里不能直接调用async方法，只是标记
            logger.warning("BrowserManager被销毁，但浏览器可能未正确关闭")


# 获取全局浏览器管理器实例
_browser_manager = BrowserManager.get_instance()


def get_browser_manager():
    """获取全局浏览器管理器"""
    return _browser_manager
# ============== 浏览器实例管理器结束 ==============


def sanitize_filename(filename: str) -> str:
    """
    清理文件名，移除或替换不允许的字符，使其符合Linux文件系统要求
    
    Args:
        filename: 原始文件名
        
    Returns:
        清理后的文件名
    """
    # Linux文件系统中不允许的字符: / < > : " | ? * \ 以及控制字符
    # 替换这些字符为下划线或移除
    import string
    
    # 不允许的字符列表
    forbidden_chars = '/<>:"|?*\\'
    
    # 替换不允许的字符为下划线
    sanitized = filename
    for char in forbidden_chars:
        sanitized = sanitized.replace(char, '_')
    
    # 移除控制字符（ASCII 0-31，除了换行符等）
    sanitized = ''.join(char for char in sanitized if ord(char) >= 32 or char in '\n\r\t')
    
    # 移除首尾空格和点号（避免隐藏文件）
    sanitized = sanitized.strip(' .')
    
    # 如果文件名为空，使用默认名称
    if not sanitized:
        sanitized = 'untitled'
    
    # 限制文件名长度（Linux通常支持255字符，但为了安全起见限制为200）
    if len(sanitized) > 200:
        # 保留扩展名
        if '.' in sanitized:
            name, ext = sanitized.rsplit('.', 1)
            sanitized = name[:200-len(ext)-1] + '.' + ext
        else:
            sanitized = sanitized[:200]
    
    return sanitized


def sanitize_file_path(file_path: str) -> str:
    """
    清理文件路径，确保路径和文件名都符合Linux文件系统要求
    
    Args:
        file_path: 原始文件路径
        
    Returns:
        清理后的文件路径
    """
    # 检查是否为绝对路径
    is_absolute = os.path.isabs(file_path)
    
    # 分离目录和文件名
    dir_path = os.path.dirname(file_path)
    filename = os.path.basename(file_path)
    
    # 清理文件名
    sanitized_filename = sanitize_filename(filename)
    
    # 如果原路径有目录部分，清理目录路径
    if dir_path:
        # 清理目录路径中的每个部分
        dir_parts = []
        for part in dir_path.split(os.sep):
            if part:  # 跳过空部分
                sanitized_part = sanitize_filename(part)
                if sanitized_part:  # 只添加非空部分
                    dir_parts.append(sanitized_part)
        
        if dir_parts:
            # 如果是绝对路径，确保以/开头
            if is_absolute:
                sanitized_dir = os.sep + os.sep.join(dir_parts)
            else:
                sanitized_dir = os.sep.join(dir_parts)
            
            # 确保目录存在
            os.makedirs(sanitized_dir, exist_ok=True)
            return os.path.join(sanitized_dir, sanitized_filename)
        else:
            # 目录部分被清理后为空，只返回文件名
            return sanitized_filename
    else:
        # 没有目录部分，直接返回清理后的文件名
        return sanitized_filename


class FileConverter:
    """文件格式转换器类"""
    
    def __init__(self, auto_cleanup=True):
        """初始化转换器"""
        self.temp_dir = tempfile.mkdtemp()
        self.auto_cleanup = auto_cleanup
        logger.info(f"临时目录创建: {self.temp_dir}")
    
    def convert_file(self, content: str, source_format: str, target_format: str, output_path: Optional[str] = None) -> str:
        """
        转换文件格式
        
        Args:
            content: 文件内容
            source_format: 源格式 (如 'md', 'html')
            target_format: 目标格式 (如 'docx', 'html', 'pdf')
            output_path: 输出文件路径
            
        Returns:
            转换后的文件路径
        """
        try:
            print(f"convert_file: {source_format}, {target_format}, {output_path}")
            # 验证输入参数
            if not content or not content.strip():
                raise ValueError("文件内容不能为空")
            
            if not source_format or not target_format:
                raise ValueError("源格式和目标格式不能为空")
            
            # 检查是否支持该转换
            if not is_conversion_supported(source_format, target_format):
                raise ValueError(f"不支持的转换: {source_format} -> {target_format}")
            
            if output_path is None:
                output_path = os.path.join(self.temp_dir, f"converted.{target_format}")
            else:
                # 如果输出路径已存在（更新场景），直接使用原始路径，不进行清理
                # 这样可以确保覆盖现有文件，而不是创建新文件
                if os.path.exists(output_path):
                    logger.info(f"输出文件已存在，将覆盖: {output_path}")
                    # 确保目录存在
                    os.makedirs(os.path.dirname(output_path), exist_ok=True)
                else:
                    # 清理文件路径，确保文件名符合Linux文件系统要求（新建文件场景）
                    output_path = sanitize_file_path(output_path)
                    logger.info(f"清理后的输出路径: {output_path}")
            
            logger.info(f"开始转换: {source_format} -> {target_format}")
            logger.info(f"输出路径: {output_path}")
            
            # 根据源格式和目标格式选择转换方法
            if source_format == 'md' and target_format == 'html':
                result = self._md_to_html(content, output_path)
            elif source_format == 'md' and target_format == 'docx':
                result = self._md_to_docx(content, output_path)
            elif source_format == 'md' and target_format == 'pdf':
                result = self._md_to_pdf(content, output_path)
            elif source_format == 'html' and target_format == 'pdf':
                result = self._html_to_pdf(content, output_path)
            elif source_format == 'html' and target_format == 'docx':
                result = self._html_to_docx(content, output_path)
            else:
                raise ValueError(f"不支持的转换: {source_format} -> {target_format}")
            
            # 验证输出文件是否存在
            if not os.path.exists(result):
                raise RuntimeError(f"转换失败：输出文件不存在 {result}")
            
            file_size = os.path.getsize(result)
            logger.info(f"转换成功: {result} (大小: {file_size} bytes)")
            return result
                
        except Exception as e:
            logger.error(f"文件转换失败: {e}")
            logger.error(f"转换参数: {source_format} -> {target_format}")
            raise
    
    def _md_to_html(self, content: str, output_path: str) -> str:
        """Markdown转HTML"""
        try:
            logger.info("开始Markdown转HTML转换")
            
            # 配置markdown扩展
            extensions = [
                'tables',
                'codehilite',
                'fenced_code',
                'toc',
                'nl2br',
                'attr_list'
            ]
            
            # 创建markdown实例
            md = markdown.Markdown(extensions=extensions)
            
            # 转换为HTML
            html_content = md.convert(content)
            
            if not html_content:
                raise ValueError("Markdown转换后HTML内容为空")
            
            # 添加完整的HTML结构
            full_html = f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>融梦AI</title>
    <style>
        body {{
            font-family: 'Alibaba PuHuiTi 3.0', 'Microsoft YaHei', 'SimHei', 'STHeiti', -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Roboto', 'Helvetica Neue', Arial, 'Noto Color Emoji', 'Segoe UI Emoji', 'Apple Color Emoji', sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #fff;
        }}
        h1, h2, h3, h4, h5, h6 {{
            color: #2c3e50;
            margin-top: 24px;
            margin-bottom: 16px;
            font-weight: 600;
            line-height: 1.25;
        }}
        h1 {{
            font-size: 2em;
            border-bottom: 1px solid #eaecef;
            padding-bottom: 0.3em;
        }}
        h2 {{
            font-size: 1.5em;
            border-bottom: 1px solid #eaecef;
            padding-bottom: 0.3em;
        }}
        table {{
            border-collapse: collapse;
            border-spacing: 0;
            width: 100%;
            margin-bottom: 16px;
        }}
        table th, table td {{
            padding: 6px 13px;
            border: 1px solid #dfe2e5;
        }}
        table th {{
            font-weight: 600;
            background-color: #f6f8fa;
        }}
        table tr:nth-child(2n) {{
            background-color: #f6f8fa;
        }}
        code {{
            padding: 0.2em 0.4em;
            margin: 0;
            font-size: 85%;
            background-color: rgba(27, 31, 35, 0.05);
            border-radius: 3px;
            font-family: 'SFMono-Regular', Consolas, 'Liberation Mono', Menlo, monospace;
        }}
        pre {{
            padding: 16px;
            overflow: auto;
            font-size: 85%;
            line-height: 1.45;
            background-color: #f6f8fa;
            border-radius: 3px;
            margin-bottom: 16px;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""
            
            # 保存HTML文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_html)
            
            logger.info(f"Markdown转HTML成功: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Markdown转HTML失败: {e}")
            raise
    
    def _md_to_docx(self, content: str, output_path: str) -> str:
        """Markdown转DOCX"""
        try:
            logger.info("开始Markdown转DOCX转换")
            logger.info(f"输入内容长度: {len(content)}")
            logger.info(f"输出路径: {output_path}")
            
            # 先转换为HTML
            logger.info("步骤1: Markdown转HTML")
            html_content = self._md_to_html(content, os.path.join(self.temp_dir, "temp.html"))
            logger.info(f"HTML文件生成: {html_content}")
            
            # 读取HTML内容
            with open(html_content, 'r', encoding='utf-8') as f:
                html_text = f.read()
            
            logger.info(f"HTML内容长度: {len(html_text)}")
            logger.info(f"HTML内容前200字符: {html_text[:200]}")
            
            if not html_text:
                raise ValueError("HTML内容为空")
            
            # 使用HTML转DOCX的方法，这样更稳定
            logger.info("步骤2: HTML转DOCX")
            result = self._html_to_docx(html_text, output_path)
            
            # 验证生成的DOCX文件
            if os.path.exists(result):
                file_size = os.path.getsize(result)
                logger.info(f"DOCX文件生成成功: {result}, 大小: {file_size} bytes")
                
                # 检查文件头
                with open(result, 'rb') as f:
                    header = f.read(4)
                    logger.info(f"DOCX文件头: {header}")
                    if header == b'PK\x03\x04':
                        logger.info("✅ DOCX文件格式正确")
                    else:
                        logger.warning(f"⚠️ DOCX文件格式可能有问题，文件头: {header}")
            else:
                logger.error("❌ DOCX文件生成失败")
            
            logger.info(f"Markdown转DOCX成功: {result}")
            return result
                
        except Exception as e:
            logger.error(f"Markdown转DOCX失败: {e}")
            import traceback
            logger.error(f"错误堆栈: {traceback.format_exc()}")
            raise
    
    def _md_to_pdf(self, content: str, output_path: str) -> str:
        """Markdown转PDF"""
        try:
            # 先转换为HTML
            html_path = self._md_to_html(content, os.path.join(self.temp_dir, "temp.html"))
            
            # 读取HTML内容
            with open(html_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # 使用改进的HTML转PDF方法（支持图片）
            return self._html_to_pdf(html_content, output_path)
            
        except Exception as e:
            logger.error(f"Markdown转PDF失败: {e}")
            raise
    
    def _html_to_pdf(self, content: str, output_path: str) -> str:
        """HTML转PDF"""
        try:
            # 优先使用Playwright（更好的图片支持和渲染效果）
            if PLAYWRIGHT_AVAILABLE:
                try:
                    logger.info("优先使用Playwright生成PDF")
                    result = self._html_to_pdf_playwright(content, output_path)
                    # 验证文件大小，确保不是空文件
                    if os.path.exists(result) and os.path.getsize(result) > 0:
                        file_size = os.path.getsize(result)
                        logger.info(f"HTML转PDF成功（Playwright）: {result} (大小: {file_size} bytes)")
                        return result
                    else:
                        logger.warning(f"Playwright生成的PDF文件为空或不存在，尝试weasyprint")
                        # 如果文件存在但为空，删除它
                        if os.path.exists(result):
                            os.remove(result)
                except Exception as e:
                    logger.warning(f"Playwright方法失败，尝试weasyprint: {e}")
                    # 如果playwright失败，确保输出文件不存在
                    if os.path.exists(output_path):
                        os.remove(output_path)
            else:
                logger.info("Playwright不可用，将使用weasyprint")
            
            # 备用方法：使用weasyprint（中文字体支持）
            try:
                logger.info("使用weasyprint生成PDF")
                weasyprint.HTML(string=content).write_pdf(output_path)
                
                # 验证文件大小，确保不是空文件
                if os.path.exists(output_path) and os.path.getsize(output_path) > 1000:
                    file_size = os.path.getsize(output_path)
                    logger.info(f"HTML转PDF成功（weasyprint）: {output_path} (大小: {file_size} bytes)")
                    return output_path
                else:
                    logger.warning(f"Weasyprint生成的PDF文件过小或不存在")
                    if os.path.exists(output_path):
                        os.remove(output_path)
            except Exception as e:
                logger.warning(f"Weasyprint方法也失败: {e}")
                if os.path.exists(output_path):
                    os.remove(output_path)
            
            # 如果两种方法都失败，抛出异常
            raise RuntimeError("所有PDF生成方法都失败了")
            
        except Exception as e:
            logger.error(f"HTML转PDF失败: {e}")
            # 确保失败时删除可能存在的空文件
            if os.path.exists(output_path) and os.path.getsize(output_path) == 0:
                try:
                    os.remove(output_path)
                except:
                    pass
            raise
    
    def _preprocess_html_images(self, content: str) -> str:
        """预处理HTML内容，将网络图片转换为base64嵌入"""
        try:            
            # 使用BeautifulSoup解析HTML
            soup = BeautifulSoup(content, 'html.parser')
            
            # 查找所有img标签
            img_tags = soup.find_all('img')
            
            for i, img in enumerate(img_tags):
                src = img.get('src', '')
                if src and (src.startswith('http') or src.startswith('https')):
                    try:
                        # 串行下载图片，使用更长的超时时间
                        response = requests.get(src, timeout=5, headers={
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                        })
                        response.raise_for_status()
                        
                        # 获取内容类型
                        content_type = response.headers.get('content-type', 'image/jpeg')
                        
                        # 转换为base64
                        img_data = base64.b64encode(response.content).decode('utf-8')
                        data_url = f"data:{content_type};base64,{img_data}"
                        # http://127.0.0.1:8081/v1/file_tool/download/geniesession-1760320791097-6795:1760320791115-7800/tmpgxcxx4bl.png
                        # http://127.0.0.1:8081/v1/file_tool/download/geniesession-1760320791097-6795:1760320791115-7800/tmpgxcxx4bl.png
                        # 更新src属性
                        img['src'] = data_url
                        
                    except requests.exceptions.Timeout as e:
                        logger.warning(f"图片 {i+1} 下载超时: {e}")
                        # 如果超时，尝试使用占位符图片
                        svg_content = '<svg width="150" height="150" xmlns="http://www.w3.org/2000/svg"><rect width="150" height="150" fill="#f0f0f0"/><text x="75" y="75" text-anchor="middle" fill="#999" font-size="12">图片加载失败</text></svg>'
                        img['src'] = f"data:image/svg+xml;base64,{base64.b64encode(svg_content.encode()).decode()}"
                        logger.info(f"图片 {i+1} 使用占位符替代")
                        continue
                    except requests.exceptions.ConnectionError as e:
                        logger.warning(f"图片 {i+1} 连接错误: {e}")
                        # 如果连接错误，使用占位符图片
                        svg_content = '<svg width="150" height="150" xmlns="http://www.w3.org/2000/svg"><rect width="150" height="150" fill="#f0f0f0"/><text x="75" y="75" text-anchor="middle" fill="#999" font-size="12">图片加载失败</text></svg>'
                        img['src'] = f"data:image/svg+xml;base64,{base64.b64encode(svg_content.encode()).decode()}"
                        logger.info(f"图片 {i+1} 使用占位符替代")
                        continue
                    except requests.exceptions.HTTPError as e:
                        logger.warning(f"图片 {i+1} HTTP错误: {e}")
                        # 如果HTTP错误，使用占位符图片
                        svg_content = '<svg width="150" height="150" xmlns="http://www.w3.org/2000/svg"><rect width="150" height="150" fill="#f0f0f0"/><text x="75" y="75" text-anchor="middle" fill="#999" font-size="12">图片加载失败</text></svg>'
                        img['src'] = f"data:image/svg+xml;base64,{base64.b64encode(svg_content.encode()).decode()}"
                        logger.info(f"图片 {i+1} 使用占位符替代")
                        continue
                    except Exception as e:
                        logger.warning(f"图片 {i+1} 转换失败: {e}")
                        # 如果转换失败，保持原URL
                        continue
            
            # 返回处理后的HTML
            processed_html = str(soup)
            logger.info("HTML图片预处理完成")
            return processed_html
            
        except Exception as e:
            logger.error(f"HTML图片预处理失败: {e}")
            return content  # 如果预处理失败，返回原始内容

    def _html_to_pdf_playwright(self, content: str, output_path: str) -> str:
        """使用Playwright将HTML转换为PDF（每次使用独立浏览器实例，更稳定）"""
        if not PLAYWRIGHT_AVAILABLE:
            raise RuntimeError("Playwright未安装，无法使用此方法")
        
        try:            
            # 预处理HTML内容，将网络图片转换为base64
            processed_content = self._preprocess_html_images(content)
            
            # 创建临时HTML文件
            temp_html_path = os.path.join(self.temp_dir, "temp.html")
            with open(temp_html_path, 'w', encoding='utf-8') as f:
                f.write(processed_content)
            
            # 使用asyncio运行异步函数
            async def convert_async():
                playwright = None
                browser = None
                page = None
                
                try:
                    # 每次都创建独立的playwright和浏览器实例
                    playwright = await async_playwright().start()
                    browser = await playwright.chromium.launch(
                        headless=True,
                        args=[
                            '--no-sandbox',
                            '--disable-setuid-sandbox',
                            '--disable-dev-shm-usage',
                            '--font-render-hinting=none',  # 禁用字体渲染提示，提高兼容性
                            '--disable-font-subpixel-positioning',  # 禁用字体子像素定位
                        ]
                    )
                    logger.info("✅ 浏览器启动成功")
                    
                    # 创建新页面
                    page = await browser.new_page()
                    
                    # 设置超时时间
                    page.set_default_timeout(30000)  # 30秒
                    
                    # 加载本地HTML文件
                    file_url = f"file://{os.path.abspath(temp_html_path)}"
                    await page.goto(file_url, wait_until='load', timeout=30000)
                    
                    # 注入额外的CSS确保字体正确显示（包括emoji字体支持）
                    await page.add_style_tag(content="""
                        * {
                            -webkit-font-smoothing: antialiased;
                            -moz-osx-font-smoothing: grayscale;
                        }
                        body, table, th, td, h1, h2, h3, h4, h5, h6, p, div, span, li, ul, ol {
                            font-family: "Alibaba PuHuiTi 3.0", "阿里巴巴普惠体 3.0", "PingFang SC", "Microsoft YaHei", "Hiragino Sans GB", -apple-system, BlinkMacSystemFont, "Segoe UI", Arial, "Noto Color Emoji", "Segoe UI Emoji", "Apple Color Emoji", sans-serif !important;
                            color: #000 !important;
                        }
                    """)
                    
                    # 等待页面完全渲染
                    await page.wait_for_timeout(2000)
                    # 检查页面内容
                    await page.evaluate("() => document.body.innerText.substring(0, 100)")
                    # 等待所有图片加载完成（如果有图片的话）
                    has_images = await page.evaluate("() => document.querySelectorAll('img').length > 0")
                    
                    if has_images:
                        try:
                            # 等待所有图片元素出现
                            await page.wait_for_selector('img', timeout=10000)
                            
                            # 等待所有图片加载完成
                            await page.wait_for_function("""
                                () => {
                                    const images = Array.from(document.querySelectorAll('img'));
                                    if (images.length === 0) return true;
                                    return images.every(img => img.complete && img.naturalHeight !== 0);
                                }
                            """, timeout=15000)
                            logger.info("所有图片加载完成")
                        except Exception as e:
                            logger.warning(f"等待图片加载超时: {e}，继续生成PDF")
                    else:
                        logger.info("未检测到图片，跳过图片加载等待")
                    
                    # 额外等待确保渲染完成
                    await page.wait_for_timeout(1000)
                    
                    # 生成PDF
                    await page.pdf(
                        path=output_path, 
                        format='A4',
                        print_background=True,
                        prefer_css_page_size=True,
                        margin={'top': '1cm', 'right': '1cm', 'bottom': '1cm', 'left': '1cm'}
                    )
                    logger.info(f"PDF已保存至: {output_path}")
                    
                finally:
                    # 清理资源（重要：按顺序关闭）
                    if page:
                        try:
                            await page.close()
                        except Exception as e:
                            logger.warning(f"关闭页面失败: {e}")
                    
                    if browser:
                        try:
                            await browser.close()
                        except Exception as e:
                            logger.warning(f"关闭浏览器失败: {e}")
                    
                    if playwright:
                        try:
                            await playwright.stop()
                            logger.info("Playwright已停止")
                        except Exception as e:
                            logger.warning(f"停止Playwright失败: {e}")
            
            # 智能选择事件循环执行方式
            try:
                # 检查是否有运行中的事件循环
                loop = asyncio.get_running_loop()
                logger.info("检测到运行中的事件循环，使用线程池执行")
                
                # 使用线程池避免事件循环冲突
                import concurrent.futures
                
                def run_in_thread():
                    new_loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(new_loop)
                    try:
                        new_loop.run_until_complete(convert_async())
                    finally:
                        new_loop.close()
                        asyncio.set_event_loop(None)
                
                # 设置60秒超时
                with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
                    future = executor.submit(run_in_thread)
                    future.result(timeout=60)
                    
            except RuntimeError:
                # 没有运行中的事件循环，直接运行
                logger.info("没有运行中的事件循环，直接执行")
                asyncio.run(convert_async())
            
            # 验证输出文件
            if not os.path.exists(output_path):
                raise RuntimeError("PDF文件生成失败：文件不存在")
            
            file_size = os.path.getsize(output_path)
            if file_size == 0:
                raise RuntimeError("PDF文件生成失败：文件为空")
            
            logger.info(f"HTML转PDF成功（Playwright）: {output_path} (大小: {file_size} bytes)")
            return output_path
            
        except Exception as e:
            logger.error(f"Playwright HTML转PDF失败: {e}")
            import traceback
            logger.error(f"错误堆栈: {traceback.format_exc()}")
            # 确保失败时删除可能存在的空文件
            if os.path.exists(output_path) and os.path.getsize(output_path) == 0:
                try:
                    os.remove(output_path)
                except:
                    pass
            raise
    
    def _html_to_docx(self, content: str, output_path: str) -> str:
        """HTML转DOCX"""
        try:
            # 使用python-docx方法转换（支持复杂表格、colspan/rowspan、图片）
            return self._html_to_docx_fallback(content, output_path)
                
        except Exception as e:
            logger.error(f"HTML转DOCX失败: {e}")
            raise
    
    def _set_document_font(self, doc: Document, font_name: str = 'Microsoft YaHei', font_size: int = 12, 
                          chinese_font: str = 'SimSun'):
        """
        设置Word文档的默认字体（中文使用宋体，西文使用微软雅黑以支持emoji）
        
        Args:
            doc: Word文档对象
            font_name: 西文字体名称（默认：微软雅黑/Microsoft YaHei，支持emoji）
            font_size: 字体大小（磅，默认：12）
            chinese_font: 中文字体名称（默认：宋体/SimSun）
        """
        try:
            from docx.shared import Pt
            from docx.oxml.ns import qn
            
            # 获取或创建Normal样式
            styles = doc.styles
            normal_style = styles['Normal']
            
            # 设置西文字体（用于emoji等特殊字符）
            font = normal_style.font
            font.name = font_name  # 西文字体
            font.size = Pt(font_size)  # 字体大小
            
            # 设置混合字体：中文用宋体，西文/emoji用微软雅黑
            element = normal_style.element
            rPr = element.get_or_add_rPr()
            
            # 设置字体族
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                from docx.oxml import OxmlElement
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            
            rFonts.set(qn('w:eastAsia'), chinese_font)  # 东亚字体（中文）- 宋体
            rFonts.set(qn('w:ascii'), font_name)        # ASCII字符 - 微软雅黑（支持emoji）
            rFonts.set(qn('w:hAnsi'), font_name)        # 高位ANSI字符 - 微软雅黑
            rFonts.set(qn('w:cs'), font_name)           # 复杂文字 - 微软雅黑
            
        except Exception as e:
            logger.warning(f"设置文档字体失败（非致命错误）: {e}")
    
    def _html_to_docx_fallback(self, html_content: str, output_path: str) -> str:
        """HTML转DOCX备用方法"""
        try:
            logger.info("开始HTML转DOCX转换（备用方法）")
            logger.info(f"HTML内容长度: {len(html_content)}")
            logger.info(f"输出路径: {output_path}")
            
            if not html_content or not html_content.strip():
                raise ValueError("HTML内容为空")
            
            # 创建Word文档
            logger.info("创建Word文档对象")
            doc = Document()
            
            # 设置文档默认字体为微软雅黑（支持emoji）
            # 注意：微软雅黑在Windows上对emoji有更好的支持
            self._set_document_font(doc, font_name='Microsoft YaHei', font_size=12)
            
            # 解析HTML并转换为Word格式
            logger.info("解析HTML内容")
            soup = BeautifulSoup(html_content, 'html.parser')
            
            if not soup:
                raise ValueError("HTML解析失败")
            
            # 解析CSS样式规则
            logger.info("解析CSS样式规则")
            self.css_styles = self._parse_css_styles(soup)
            
            # 查找body元素，如果没有则使用整个文档
            body = soup.find('body')
            if body is None:
                body = soup
                logger.info("未找到body元素，使用整个文档")
            else:
                logger.info("找到body元素")
            
            # 按顺序处理所有子元素
            logger.info("开始处理HTML元素")
            processed_elements = 0
            for i, element in enumerate(body.contents):
                try:
                    # logger.info(f"处理元素 {i+1}: {element.name if hasattr(element, 'name') else '文本节点'}")
                    self._process_element_to_docx(doc, element)
                    processed_elements += 1
                except Exception as element_error:
                    logger.warning(f"处理HTML元素失败: {element_error}")
                    continue
            
            # logger.info(f"成功处理了 {processed_elements} 个元素")
            
            if processed_elements == 0:
                logger.warning("没有成功处理任何HTML元素，添加默认内容")
                doc.add_paragraph("文档内容")
            
            # 保存文档
            logger.info(f"保存DOCX文件到: {output_path}")
            doc.save(output_path)
            
            # 验证文件是否成功创建
            if not os.path.exists(output_path):
                raise RuntimeError("DOCX文件保存失败")
            
            file_size = os.path.getsize(output_path)
            logger.info(f"HTML转DOCX成功: {output_path} (大小: {file_size} bytes)")
            
            # 验证文件格式
            with open(output_path, 'rb') as f:
                header = f.read(4)
                logger.info(f"生成的DOCX文件头: {header}")
                if header == b'PK\x03\x04':
                    logger.info("✅ DOCX文件格式正确")
                else:
                    logger.warning(f"⚠️ DOCX文件格式可能有问题，文件头: {header}")
            
            return output_path
            
        except Exception as e:
            logger.error(f"HTML转DOCX失败: {e}")
            import traceback
            logger.error(f"错误堆栈: {traceback.format_exc()}")
            raise
    
    def _download_and_save_image(self, src: str) -> Optional[str]:
        """
        下载图片并保存到临时文件
        支持URL和base64格式
        
        Args:
            src: 图片源（URL或base64 data URI）
            
        Returns:
            临时图片文件路径，失败返回None
        """
        try:
            # 处理base64格式
            if src.startswith('data:image'):
                # 解析data URI: data:image/png;base64,iVBORw0KGgoAAAANS...
                # 支持 svg+xml 等复杂格式
                match = re.match(r'data:image/([^;]+);base64,(.+)', src)
                if match:
                    image_format = match.group(1)
                    image_data = match.group(2)
                    
                    # 处理特殊格式名称（如 svg+xml -> svg）
                    if '+' in image_format:
                        image_format = image_format.split('+')[0]
                    
                    # 解码base64
                    img_bytes = base64.b64decode(image_data)
                    
                    # 保存到临时文件
                    temp_image_path = os.path.join(self.temp_dir, f"img_{hash(src)}.{image_format}")
                    with open(temp_image_path, 'wb') as f:
                        f.write(img_bytes)
                    
                    logger.info(f"Base64图片已保存: {temp_image_path}, 大小: {len(img_bytes)} bytes")
                    return temp_image_path
            
            # 处理URL格式
            elif src.startswith('http://') or src.startswith('https://'):
                # 下载图片
                response = requests.get(src, timeout=10, headers={
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                })
                response.raise_for_status()
                
                # 从URL或Content-Type推断图片格式
                content_type = response.headers.get('content-type', '')
                if 'image/png' in content_type:
                    ext = 'png'
                elif 'image/jpeg' in content_type or 'image/jpg' in content_type:
                    ext = 'jpg'
                elif 'image/gif' in content_type:
                    ext = 'gif'
                elif 'image/webp' in content_type:
                    ext = 'webp'
                else:
                    # 从URL推断
                    ext = src.split('.')[-1].split('?')[0].lower()
                    if ext not in ['png', 'jpg', 'jpeg', 'gif', 'webp']:
                        ext = 'png'  # 默认使用png
                
                # 保存到临时文件
                temp_image_path = os.path.join(self.temp_dir, f"img_{hash(src)}.{ext}")
                with open(temp_image_path, 'wb') as f:
                    f.write(response.content)
                
                logger.info(f"URL图片已下载: {temp_image_path}, 大小: {len(response.content)} bytes")
                return temp_image_path
            
            # 处理本地文件路径
            elif os.path.exists(src):
                logger.info(f"使用本地图片: {src}")
                return src
            
            return None
            
        except Exception as e:
            logger.warning(f"下载/保存图片失败: {src}, 错误: {e}")
            return None
    
    def _add_image_to_docx(self, doc: Document, src: str, alt: str = '图片'):
        """
        将图片添加到Word文档
        
        Args:
            doc: Word文档对象
            src: 图片源（URL或base64）
            alt: 备用文本
        """
        try:
            # 下载并保存图片
            image_path = self._download_and_save_image(src)
            
            if image_path and os.path.exists(image_path):
                # 添加图片到文档，设置合适的宽度
                try:
                    # 创建一个段落来放置图片
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    
                    # 设置图片宽度为4英寸（约10cm），保持宽高比
                    run.add_picture(image_path, width=Inches(4))
                    
                    # 设置段落居中
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    logger.info(f"图片已插入文档: {image_path}")
                except Exception as insert_error:
                    logger.warning(f"插入图片失败: {insert_error}, 使用文本替代")
                    doc.add_paragraph(f"[图片: {alt}]")
            else:
                # 如果图片下载失败，添加文本占位符
                doc.add_paragraph(f"[图片加载失败: {alt}]")
                logger.warning(f"图片加载失败，使用文本占位符: {src}")
                
        except Exception as e:
            logger.error(f"添加图片到DOCX失败: {e}")
            # 添加文本占位符
            doc.add_paragraph(f"[图片: {alt}]")
    
    def _parse_css_styles(self, soup):
        """
        解析HTML中的<style>标签，提取CSS样式规则
        
        Args:
            soup: BeautifulSoup对象
            
        Returns:
            dict: CSS样式规则字典，格式为 {selector: {property: value}}
        """
        css_rules = {}
        
        try:
            # 查找所有style标签
            style_tags = soup.find_all('style')
            
            for style_tag in style_tags:
                css_content = style_tag.string
                if not css_content:
                    continue
                
                # logger.info(f"解析CSS内容，长度: {len(css_content)}")
                
                # 简单的CSS解析（支持基本的选择器和属性）
                # 匹配形如: selector { property: value; }
                # 使用正则表达式提取CSS规则
                rule_pattern = r'([^{]+)\{([^}]+)\}'
                matches = re.finditer(rule_pattern, css_content, re.MULTILINE)
                
                for match in matches:
                    selector = match.group(1).strip()
                    properties_str = match.group(2).strip()
                    
                    # 解析属性
                    properties = {}
                    prop_pattern = r'([^:;]+):([^;]+)'
                    prop_matches = re.finditer(prop_pattern, properties_str)
                    
                    for prop_match in prop_matches:
                        prop_name = prop_match.group(1).strip()
                        prop_value = prop_match.group(2).strip()
                        properties[prop_name] = prop_value
                    
                    if properties:
                        css_rules[selector] = properties
                        # logger.info(f"解析CSS规则: {selector} -> {properties}")
            
            # logger.info(f"共解析了 {len(css_rules)} 条CSS规则")
            
        except Exception as e:
            logger.warning(f"解析CSS样式失败: {e}")
        
        return css_rules
    
    def _process_element_to_docx(self, doc: Document, element):
        """递归处理HTML元素，保持正确的顺序"""
        # 过滤HTML注释
        from bs4.element import Comment
        if isinstance(element, Comment):
            # 跳过注释节点，如 <!-- [标题区] -->
            return
        
        # 检查是否为文本节点
        if hasattr(element, 'name') and element.name is None:
            # 处理文本节点
            text = self._clean_text_whitespace(str(element))
            if text:
                para = doc.add_paragraph(text)
                # 设置字体为微软雅黑（支持emoji）
                self._set_paragraph_font(para, font_name='Microsoft YaHei', font_size=12)
            return
        
        # 检查是否为标签元素
        if not hasattr(element, 'name') or element.name is None:
            return
        
        tag_name = element.name
        
        # 处理标题
        if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(tag_name[1])
            text = element.get_text().strip()
            if text:
                heading = doc.add_heading(text, level=level)
                # 设置字体为微软雅黑（支持emoji）
                self._set_paragraph_font(heading, font_name='Microsoft YaHei', font_size=16 if tag_name == 'h1' else 14)
                # 检查并应用居中样式
                alignment = self._get_text_alignment(element)
                if alignment:
                    heading.alignment = alignment
        
        # 处理段落
        elif tag_name == 'p':
            # 检查段落是否包含图片
            has_images = element.find_all(['img', 'svg'])
            if has_images:
                # 如果段落包含图片，需要分别处理文本和图片
                if hasattr(element, 'contents'):
                    # 获取段落的对齐方式
                    alignment = self._get_text_alignment(element)
                    
                    for child in element.contents:
                        # 过滤HTML注释
                        from bs4.element import Comment
                        if isinstance(child, Comment):
                            continue
                        
                        if hasattr(child, 'name') and child.name == 'img':
                            src = child.get('src', '') if hasattr(child, 'get') else ''
                            alt = child.get('alt', '图片') if hasattr(child, 'get') else '图片'
                            if src:
                                # 真正插入图片
                                self._add_image_to_docx(doc, src, alt)
                        elif not hasattr(child, 'name') or child.name is None:
                            # 处理文本节点
                            text = self._clean_text_whitespace(str(child))
                            if text:
                                para = doc.add_paragraph(text)
                                # 设置字体为微软雅黑（支持emoji）
                                self._set_paragraph_font(para, font_name='Microsoft YaHei', font_size=12)
                                # 应用对齐方式
                                if alignment:
                                    para.alignment = alignment
            else:
                # 普通段落
                text = self._clean_text_whitespace(element.get_text())
                if text:
                    para = doc.add_paragraph(text)
                    # 设置字体为微软雅黑（支持emoji）
                    self._set_paragraph_font(para, font_name='Microsoft YaHei', font_size=12)
                    # 检查并应用居中样式
                    alignment = self._get_text_alignment(element)
                    if alignment:
                        para.alignment = alignment
        
        # 处理独立的图片
        elif tag_name == 'img':
            src = element.get('src', '') if hasattr(element, 'get') else ''
            alt = element.get('alt', '图片') if hasattr(element, 'get') else '图片'
            if src:
                # 真正插入图片
                self._add_image_to_docx(doc, src, alt)
            else:
                para = doc.add_paragraph(f"[{alt}]")
                # 设置字体为微软雅黑（支持emoji）
                self._set_paragraph_font(para, font_name='Microsoft YaHei', font_size=12)
        
        # 处理表格
        elif tag_name == 'table':
            self._add_table_to_docx(doc, element)
        
        # 处理列表
        elif tag_name in ['ul', 'ol']:
            for li in element.find_all('li', recursive=False):
                text = self._clean_text_whitespace(li.get_text())
                if text:
                    if tag_name == 'ul':
                        para = doc.add_paragraph(text, style='List Bullet')
                    else:  # ol
                        para = doc.add_paragraph(text, style='List Number')
                    # 设置字体为微软雅黑（支持emoji）
                    self._set_paragraph_font(para, font_name='Microsoft YaHei', font_size=12)
        
        # 处理div和其他容器元素
        elif tag_name in ['div', 'section', 'article', 'main', 'body']:
            # 递归处理子元素，保持顺序
            if hasattr(element, 'contents'):
                for child in element.contents:
                    self._process_element_to_docx(doc, child)
        
        # 处理其他元素
        else:
            # 对于其他元素，提取文本内容
            text = self._clean_text_whitespace(element.get_text())
            if text:
                para = doc.add_paragraph(text)
                # 设置字体为微软雅黑（支持emoji）
                self._set_paragraph_font(para, font_name='Microsoft YaHei', font_size=12)

    def _get_text_alignment(self, element):
        """
        从HTML元素的style属性或CSS规则中提取text-align对齐方式
        
        Args:
            element: BeautifulSoup元素对象
            
        Returns:
            WD_ALIGN_PARAGRAPH对齐常量，如果没有找到对齐样式则返回None
        """
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            align_value = None
            
            # 1. 首先检查inline style属性（优先级最高）
            style = element.get('style', '') if hasattr(element, 'get') else ''
            
            if style:
                # 解析style字符串，提取text-align属性
                style_lower = style.lower()
                
                if 'text-align' in style_lower:
                    match = re.search(r'text-align\s*:\s*(\w+)', style_lower)
                    if match:
                        align_value = match.group(1).strip()
            
            # 2. 如果没有inline style，检查CSS规则
            if not align_value and hasattr(self, 'css_styles') and self.css_styles:
                tag_name = element.name if hasattr(element, 'name') else None
                
                if tag_name:
                    # 检查标签选择器（如 h1, p, div）
                    if tag_name in self.css_styles:
                        css_props = self.css_styles[tag_name]
                        if 'text-align' in css_props:
                            align_value = css_props['text-align'].strip()
                            logger.info(f"从CSS规则获取对齐方式: {tag_name} -> {align_value}")
                    
                    # 检查类选择器（如 .center）
                    class_names = element.get('class', []) if hasattr(element, 'get') else []
                    for class_name in class_names:
                        class_selector = f'.{class_name}'
                        if class_selector in self.css_styles:
                            css_props = self.css_styles[class_selector]
                            if 'text-align' in css_props:
                                align_value = css_props['text-align'].strip()
                                logger.info(f"从CSS类规则获取对齐方式: {class_selector} -> {align_value}")
                                break
            
            # 3. 根据对齐值返回相应的常量
            if align_value:
                if align_value == 'center':
                    logger.info(f"应用居中样式: {element.name}")
                    return WD_ALIGN_PARAGRAPH.CENTER
                elif align_value == 'right':
                    return WD_ALIGN_PARAGRAPH.RIGHT
                elif align_value == 'left':
                    return WD_ALIGN_PARAGRAPH.LEFT
                elif align_value == 'justify':
                    return WD_ALIGN_PARAGRAPH.JUSTIFY
            
            return None
            
        except Exception as e:
            logger.warning(f"解析文本对齐样式失败: {e}")
            return None
    
    def _set_paragraph_font(self, paragraph, font_name: str = 'Microsoft YaHei', font_size: int = 12,
                          chinese_font: str = 'SimSun'):
        """
        为段落设置字体（中文使用宋体，西文使用微软雅黑以支持emoji）
        
        Args:
            paragraph: Word段落对象
            font_name: 西文字体名称（默认：微软雅黑/Microsoft YaHei，支持emoji）
            font_size: 字体大小（磅，默认：12）
            chinese_font: 中文字体名称（默认：宋体/SimSun）
        """
        try:
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            # 为段落中的所有run设置字体
            for run in paragraph.runs:
                run.font.name = font_name  # 西文字体
                run.font.size = Pt(font_size)
                
                # 设置混合字体：中文用宋体，西文/emoji用微软雅黑
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                
                rFonts.set(qn('w:eastAsia'), chinese_font)  # 东亚字体（中文）- 宋体
                rFonts.set(qn('w:ascii'), font_name)        # ASCII字符 - 微软雅黑
                rFonts.set(qn('w:hAnsi'), font_name)        # 高位ANSI字符 - 微软雅黑
                rFonts.set(qn('w:cs'), font_name)           # 复杂文字 - 微软雅黑
        except Exception as e:
            logger.warning(f"设置段落字体失败: {e}")
    
    def _clean_text_whitespace(self, text: str) -> str:
        """
        清理文本中的多余空白字符
        
        Args:
            text: 原始文本
            
        Returns:
            清理后的文本
        """
        if not text:
            return text
        
        # 1. 将多个连续空格替换为单个空格
        text = re.sub(r' +', ' ', text)
        
        # 2. 将制表符替换为空格
        text = text.replace('\t', ' ')
        
        # 3. 去除换行符前后的空格
        text = re.sub(r' *\n *', '\n', text)
        
        # 4. 将多个连续换行符替换为单个换行符
        text = re.sub(r'\n+', '\n', text)
        
        # 5. 去除首尾空白
        text = text.strip()
        
        return text
    
    def _add_table_to_docx(self, doc: Document, table_soup):
        """将HTML表格添加到Word文档，支持表格单元格中的图片、colspan和rowspan"""
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            # 获取表格行（包括thead和tbody中的行）
            rows = table_soup.find_all('tr')
            if not rows:
                return
            
            # 计算实际需要的列数（考虑colspan）
            max_cols = 0
            for row in rows:
                cells = row.find_all(['td', 'th'])
                col_count = sum(int(cell.get('colspan', 1)) for cell in cells)
                max_cols = max(max_cols, col_count)
            
            if max_cols == 0:
                return
            
            # 计算实际需要的行数（考虑rowspan）
            row_count = len(rows)
            
            logger.info(f"创建表格: {row_count}行 x {max_cols}列")
            
            # 创建表格
            table = doc.add_table(rows=row_count, cols=max_cols)
            table.style = 'Table Grid'
            
            # 跟踪已合并的单元格
            merged_cells = set()
            
            # 填充表格数据
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                col_idx = 0
                
                for cell in cells:
                    # 跳过已被rowspan占用的单元格
                    while (i, col_idx) in merged_cells:
                        col_idx += 1
                    
                    if col_idx >= max_cols:
                        break
                    
                    # 获取colspan和rowspan
                    colspan = int(cell.get('colspan', 1))
                    rowspan = int(cell.get('rowspan', 1))
                    
                    # 获取起始单元格
                    try:
                        start_cell = table.cell(i, col_idx)
                        
                        # 如果有合并，处理合并
                        if colspan > 1 or rowspan > 1:
                            end_row = min(i + rowspan - 1, row_count - 1)
                            end_col = min(col_idx + colspan - 1, max_cols - 1)
                            end_cell = table.cell(end_row, end_col)
                            
                            # 合并单元格
                            if start_cell != end_cell:
                                start_cell.merge(end_cell)
                            
                            # 标记被占用的单元格
                            for r in range(i, end_row + 1):
                                for c in range(col_idx, end_col + 1):
                                    if r != i or c != col_idx:
                                        merged_cells.add((r, c))
                        
                        # 填充单元格内容
                        docx_cell = start_cell
                        
                        # 检查单元格是否包含图片
                        img_tags = cell.find_all('img')
                        if img_tags:
                            # 清空单元格
                            docx_cell.text = ''
                            
                            # 处理单元格内容（文本和图片混合）
                            for content in cell.contents:
                                # 过滤HTML注释
                                from bs4.element import Comment
                                if isinstance(content, Comment):
                                    continue
                                
                                if hasattr(content, 'name'):
                                    if content.name == 'img':
                                        # 处理图片
                                        src = content.get('src', '')
                                        alt = content.get('alt', '图片')
                                        if src:
                                            # 下载并保存图片
                                            image_path = self._download_and_save_image(src)
                                            if image_path and os.path.exists(image_path):
                                                try:
                                                    # 在单元格中添加图片
                                                    paragraph = docx_cell.paragraphs[0] if docx_cell.paragraphs else docx_cell.add_paragraph()
                                                    run = paragraph.add_run()
                                                    # 表格中的图片使用较小的尺寸
                                                    run.add_picture(image_path, width=Inches(1.2))
                                                    logger.info(f"表格图片已插入: {image_path}")
                                                except Exception as img_error:
                                                    logger.warning(f"表格图片插入失败: {img_error}")
                                                    if docx_cell.text:
                                                        docx_cell.text += f"\n[图片: {alt}]"
                                                    else:
                                                        docx_cell.text = f"[图片: {alt}]"
                                            else:
                                                if docx_cell.text:
                                                    docx_cell.text += f"\n[图片: {alt}]"
                                                else:
                                                    docx_cell.text = f"[图片: {alt}]"
                                    elif content.name == 'br':
                                        # 换行
                                        if docx_cell.text:
                                            docx_cell.text += '\n'
                                    else:
                                        # 处理其他HTML标签的文本
                                        text = self._clean_text_whitespace(content.get_text())
                                        if text:
                                            if docx_cell.text:
                                                docx_cell.text += text
                                            else:
                                                docx_cell.text = text
                                elif str(content).strip():
                                    # 处理纯文本节点
                                    text = self._clean_text_whitespace(str(content))
                                    if text and text not in ['\n', '\t']:
                                        if docx_cell.text:
                                            docx_cell.text += text
                                        else:
                                            docx_cell.text = text
                        else:
                            # 普通单元格，只有文本
                            cell_text = self._clean_text_whitespace(cell.get_text())
                            docx_cell.text = cell_text
                        
                        # 检查是否为表头单元格（th标签），如果是则应用表头样式
                        # 注意：必须在文本内容填充完成后应用样式
                        if cell.name == 'th':
                            self._apply_table_header_style(docx_cell)
                        
                        col_idx += colspan
                        
                    except Exception as cell_error:
                        logger.warning(f"处理单元格失败 ({i},{col_idx}): {cell_error}")
                        col_idx += 1
                        continue
            
            # 设置表格字体为微软雅黑（在填充内容后，支持emoji）
            try:
                self._set_table_font(table, font_name='Microsoft YaHei', font_size=9)
            except Exception as font_error:
                logger.warning(f"设置表格字体失败: {font_error}")
            
            # 为表格添加边框
            try:
                for row in table.rows:
                    for cell in row.cells:
                        self._set_cell_border(
                            cell,
                            top={"sz": 4, "val": "single", "color": "000000"},
                            bottom={"sz": 4, "val": "single", "color": "000000"},
                            left={"sz": 4, "val": "single", "color": "000000"},
                            right={"sz": 4, "val": "single", "color": "000000"}
                        )
                # logger.info(f"已为表格添加边框")
            except Exception as border_error:
                logger.warning(f"添加表格边框失败: {border_error}")
            
            # 添加换行
            doc.add_paragraph()
            
        except Exception as e:
            logger.error(f"添加表格到DOCX失败: {e}")
            import traceback
            logger.error(f"错误堆栈: {traceback.format_exc()}")
    
    def _apply_table_header_style(self, cell):
        """
        为表头单元格应用样式：加粗、灰色背景、居中对齐
        
        Args:
            cell: Word表格单元格对象
        """
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # 设置灰色背景 (#F0F0F0)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F0F0F0')
            cell._element.get_or_add_tcPr().append(shading_elm)
            
            # 设置单元格中所有段落为加粗和居中
            for paragraph in cell.paragraphs:
                # 居中对齐
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 加粗所有文本
                for run in paragraph.runs:
                    run.font.bold = True
            
            # logger.debug(f"已应用表头样式：加粗、灰色背景、居中")
            
        except Exception as e:
            logger.warning(f"应用表头样式失败: {e}")
    
    def _set_table_font(self, table, font_name: str = 'Microsoft YaHei', font_size: int = 9,
                       chinese_font: str = 'SimSun'):
        """
        设置表格的字体（中文使用宋体，西文使用微软雅黑以支持emoji）
        
        Args:
            table: 表格对象
            font_name: 西文字体名称（默认：微软雅黑/Microsoft YaHei，支持emoji）
            font_size: 字体大小（磅，默认：9）
            chinese_font: 中文字体名称（默认：宋体/SimSun）
        """
        try:
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            for row in table.rows:
                for cell in row.cells:
                    # 设置单元格中所有段落的字体
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            # 设置run的西文字体
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
                            
                            # 设置混合字体：中文用宋体，西文/emoji用微软雅黑
                            rPr = run._element.get_or_add_rPr()
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is None:
                                rFonts = OxmlElement('w:rFonts')
                                rPr.append(rFonts)
                            
                            rFonts.set(qn('w:eastAsia'), chinese_font)  # 东亚字体（中文）- 宋体
                            rFonts.set(qn('w:ascii'), font_name)        # ASCII字符 - 微软雅黑
                            rFonts.set(qn('w:hAnsi'), font_name)        # 高位ANSI字符 - 微软雅黑
                            rFonts.set(qn('w:cs'), font_name)           # 复杂文字 - 微软雅黑
            
        except Exception as e:
            logger.warning(f"设置表格字体失败（非致命错误）: {e}")
    
    def _set_cell_border(self, cell, **kwargs):
        """
        设置单元格边框
        
        参数:
            top, bottom, left, right: 边框样式字典
        """
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            
            # 创建边框元素
            tcBorders = OxmlElement('w:tcBorders')
            
            for edge in ('top', 'left', 'bottom', 'right'):
                if edge in kwargs:
                    edge_data = kwargs.get(edge)
                    edge_el = OxmlElement(f'w:{edge}')
                    
                    for key in ['sz', 'val', 'color']:
                        if key in edge_data:
                            edge_el.set(qn(f'w:{key}'), str(edge_data[key]))
                    
                    tcBorders.append(edge_el)
            
            tcPr.append(tcBorders)
        except Exception as e:
            logger.warning(f"设置单元格边框失败: {e}")
    
    def cleanup(self):
        """清理临时文件"""
        try:
            import shutil
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
                logger.info(f"临时目录已清理: {self.temp_dir}")
        except Exception as e:
            logger.error(f"清理临时目录失败: {e}")
    
    def __del__(self):
        """析构函数，自动清理临时文件"""
        if self.auto_cleanup:
            self.cleanup()


def convert_file_format(content: str, source_format: str, target_format: str, output_path: Optional[str] = None) -> str:
    """
    便捷函数：转换文件格式
    
    Args:
        content: 文件内容
        source_format: 源格式 (如 'md', 'html')
        target_format: 目标格式 (如 'docx', 'html', 'pdf')
        output_path: 输出文件路径
        
    Returns:
        转换后的文件路径
    """
    print(f"convert_file_format: {source_format}, {target_format}, {output_path}")
    converter = FileConverter(auto_cleanup=False)  # 禁用自动清理
    try:
        result_path = converter.convert_file(content, source_format, target_format, output_path)
        return result_path
    except Exception as e:
        # 只有在出错时才清理
        converter.cleanup()
        raise


# 支持的转换格式映射
SUPPORTED_CONVERSIONS = {
    'md': ['html', 'docx', 'pdf'],
    'html': ['pdf', 'docx'],
}

def is_conversion_supported(source_format: str, target_format: str) -> bool:
    """检查是否支持指定的格式转换"""
    if source_format in SUPPORTED_CONVERSIONS:
        return target_format in SUPPORTED_CONVERSIONS[source_format]
    return False


if __name__ == "__main__":
    # 测试代码
    test_markdown = """
# 测试文档

这是一个**测试**文档，用于验证文件转换功能。

## 功能列表

- Markdown转HTML
- Markdown转DOCX  
- Markdown转PDF
- HTML转PDF
- HTML转DOCX

### 代码示例

```python
def hello_world():
    print("Hello, World!")
```

| 功能 | 状态 |
|------|------|
| HTML | ✅ |
| DOCX | ✅ |
| PDF | ✅ |
"""
    
    # 测试HTML转PDF（包含图片）
    test_html_with_images = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>图片测试文档</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            margin: 20px;
            line-height: 1.6;
        }
        .image-container {
            text-align: center;
            margin: 20px 0;
        }
        img {
            max-width: 100%;
            height: auto;
            border: 1px solid #ddd;
            border-radius: 4px;
        }
    </style>
</head>
<body>
    <h1>图片转换测试</h1>
    <p>这个文档用于测试HTML转PDF时的图片处理功能。</p>
    
    <div class="image-container">
        <h2>测试图片1：在线图片</h2>
        <img src="https://via.placeholder.com/300x200/0066CC/FFFFFF?text=Test+Image+1" alt="测试图片1">
        <p>这是一个在线图片，用于测试网络图片的加载和转换。</p>
    </div>
    
    <div class="image-container">
        <h2>测试图片2：SVG图片</h2>
        <svg width="300" height="200" xmlns="http://www.w3.org/2000/svg">
            <rect width="300" height="200" fill="#4CAF50"/>
            <circle cx="150" cy="100" r="50" fill="#FFFFFF"/>
            <text x="150" y="110" text-anchor="middle" fill="#4CAF50" font-size="16" font-family="Arial">SVG Test</text>
        </svg>
        <p>这是一个SVG图片，用于测试矢量图形的转换。</p>
    </div>
    
    <h2>总结</h2>
    <p>如果PDF中能正确显示上述图片，说明图片转换功能工作正常。</p>
</body>
</html>
"""
    
    print("开始测试文件转换功能...")
    
    try:
        # 测试Markdown转HTML
        html_path = convert_file_format(test_markdown, 'md', 'html')
        print(f"Markdown转HTML成功: {html_path}")
        
        # 测试Markdown转DOCX
        docx_path = convert_file_format(test_markdown, 'md', 'docx')
        print(f"Markdown转DOCX成功: {docx_path}")
        
        # 测试Markdown转PDF
        pdf_path = convert_file_format(test_markdown, 'md', 'pdf')
        print(f"Markdown转PDF成功: {pdf_path}")
        
        # 测试HTML转PDF（包含图片）
        print("\n开始测试HTML转PDF（包含图片）...")
        html_pdf_path = convert_file_format(test_html_with_images, 'html', 'pdf')
        print(f"HTML转PDF成功: {html_pdf_path}")
        
    except Exception as e:
        print(f"转换测试失败: {e}")
        import traceback
        traceback.print_exc()
